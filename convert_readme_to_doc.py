import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup

def convert_md_to_html(md_content):
    """Convert markdown content to HTML."""
    return markdown.markdown(md_content, extensions=['tables', 'fenced_code'])

def clean_html(html_content):
    """Clean HTML content."""
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup

def add_md_content_to_doc(doc, md_file_path):
    """Add markdown content to Word document."""
    # Get the filename without extension for the section title
    file_name = os.path.basename(md_file_path)
    file_name_without_ext = os.path.splitext(file_name)[0]
    
    # Replace hyphens with spaces and capitalize for better readability
    section_title = file_name_without_ext.replace('-', ' ').title()
    
    # Add section title
    heading = doc.add_heading(section_title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add file path as reference
    file_ref = doc.add_paragraph(f"Source: {md_file_path}")
    file_ref.italic = True
    
    # Add a separator
    doc.add_paragraph("---")
    
    # Read markdown content
    try:
        with open(md_file_path, 'r', encoding='utf-8') as md_file:
            md_content = md_file.read()
        
        # Convert markdown to HTML
        html_content = convert_md_to_html(md_content)
        soup = clean_html(html_content)
        
        # Process HTML elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'pre', 'code', 'table']):
            if element.name.startswith('h'):
                # Extract heading level
                level = int(element.name[1]) + 1  # +1 because we used level 1 for the file name
                heading = doc.add_heading(element.get_text(), level=min(level, 9))
            
            elif element.name == 'p':
                p = doc.add_paragraph(element.get_text())
            
            elif element.name == 'ul':
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(li.get_text(), style='List Bullet')
            
            elif element.name == 'ol':
                for li in element.find_all('li', recursive=False):
                    p = doc.add_paragraph(li.get_text(), style='List Number')
            
            elif element.name == 'pre' or element.name == 'code':
                code_text = element.get_text()
                p = doc.add_paragraph(code_text)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            
            # Tables are complex to handle in python-docx, so we'll convert them to simple text
            elif element.name == 'table':
                doc.add_paragraph("Table content:")
                for row in element.find_all('tr'):
                    cells = [cell.get_text().strip() for cell in row.find_all(['th', 'td'])]
                    doc.add_paragraph(" | ".join(cells))
        
        # Add a page break after each file
        doc.add_page_break()
        
        return True
    except Exception as e:
        doc.add_paragraph(f"Error processing file: {str(e)}")
        doc.add_page_break()
        return False

def main():
    # Create a new Document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = "KAI Documentation"
    doc.core_properties.author = "KAI Team"
    
    # Add a title page
    title = doc.add_heading("KAI Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a subtitle
    subtitle = doc.add_paragraph("Compiled from /readme/ folder")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    import datetime
    date = doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a page break after the title page
    doc.add_page_break()
    
    # Add table of contents placeholder
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("(Table of contents will be generated when opened in Word)")
    
    # Add a page break after the TOC
    doc.add_page_break()
    
    # Get all markdown files from the readme directory
    readme_dir = "readme"
    md_files = []
    
    for root, dirs, files in os.walk(readme_dir):
        for file in files:
            if file.endswith(".md"):
                md_files.append(os.path.join(root, file))
    
    # Sort files alphabetically
    md_files.sort()
    
    # Process each markdown file
    successful_files = 0
    failed_files = 0
    
    for md_file in md_files:
        print(f"Processing: {md_file}")
        success = add_md_content_to_doc(doc, md_file)
        if success:
            successful_files += 1
        else:
            failed_files += 1
    
    # Add a summary at the end
    summary = doc.add_heading("Processing Summary", level=1)
    doc.add_paragraph(f"Total files processed: {len(md_files)}")
    doc.add_paragraph(f"Successfully processed: {successful_files}")
    doc.add_paragraph(f"Failed to process: {failed_files}")
    
    # Save the document
    output_file = "KAI_Documentation.docx"
    doc.save(output_file)
    print(f"Document saved as {output_file}")

if __name__ == "__main__":
    main()
